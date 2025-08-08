from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
import io

def export_as_docx(result):
    doc = Document()
    doc.add_heading(result["title"], 0)
    for section in result["sections"]:
        doc.add_heading(section["heading"], level=1)
        doc.add_paragraph(section["content"])
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def export_as_pdf(result):
    pdf_io = io.BytesIO()
    pdf = SimpleDocTemplate(pdf_io, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [Paragraph(f"<b>{result['title']}</b>", styles["Title"]), Spacer(1, 12)]

    for section in result["sections"]:
        elements.append(Paragraph(f"<b>{section['heading']}</b>", styles["Heading2"]))
        elements.append(Paragraph(section["content"], styles["Normal"]))
        elements.append(Spacer(1, 12))

    pdf.build(elements)
    pdf_io.seek(0)
    return pdf_io
