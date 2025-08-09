# services/file_utils.py
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
import io

def export_as_docx(result: dict) -> bytes:
    doc = Document()
    doc.add_heading(result.get("title", "Document"), level=0)
    for section in result.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=1)
        doc.add_paragraph(section.get("content", ""))
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem.read()

def export_as_pdf(result: dict) -> bytes:
    pdf_io = io.BytesIO()
    pdf = SimpleDocTemplate(pdf_io, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []
    elements.append(Paragraph(result.get("title", ""), styles["Title"]))
    elements.append(Spacer(1, 12))
    for section in result.get("sections", []):
        elements.append(Paragraph(section.get("heading", ""), styles["Heading2"]))
        elements.append(Paragraph(section.get("content", ""), styles["BodyText"]))
        elements.append(Spacer(1, 12))
    pdf.build(elements)
    pdf_io.seek(0)
    return pdf_io.read()
